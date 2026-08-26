# -*- coding: utf-8 -*-
"""McKinsey Classic Executive Theme DOCX 보고서 생성기 (PART 2 흐름 재구성 완료본)"""
import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


class DocxGenerator:
    @classmethod
    def generate(cls, data, output_docx_path):
        return cls.generate_report(data, output_docx_path)
    """맥킨지 클래식 이그제큐티브 워드 보고서 생성기"""

    @staticmethod
    def generate_report(data, output_docx_path):
        doc = docx.Document()
        
        # 여백 설정 (1인치)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            
        site = data['site']
        demo = data['demographics']
        comm = data['commercial']
        score = data.get('score', data.get('scores', {}))
        fin = data['financials']
        inv = fin['investment']
        
        # 표지 헤더 및 제목
        p_top = doc.add_paragraph("MYPARK SCREEN PARK GOLF  |  EXECUTIVE SITE SELECTION REPORT")
        p_top.runs[0].font.name = 'Malgun Gothic'
        p_top.runs[0].font.size = Pt(9.5)
        p_top.runs[0].font.bold = True
        p_top.runs[0].font.color.rgb = RGBColor(0, 43, 73)
        
        h_main = doc.add_heading("스크린 파크골프(마이파크) 출점 타당성 분석 보고서", level=0)
        h_main.runs[0].font.name = 'Malgun Gothic'
        h_main.runs[0].font.color.rgb = RGBColor(0, 43, 73)
        
        doc.add_paragraph("10타석 120평 플래그십 표준 모델  |  상권 분석 및 투자 타당성 평가").paragraph_format.space_after = Pt(14)
        
        # 메타 정보 요약 테이블
        meta_table = doc.add_table(rows=5, cols=2)
        meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        meta_data = [
            ("대상 사업지", f"{site['full_address']}" + (f" (특이사항: {site['special_notes']})" if site.get('special_notes') else "")),
            ("표준 출점 모델", f"{site['rooms']}타석 ({site['area_pyeong']}평형 플래그십 모델)"),
            ("상권 분석 범위", f"{site['sido']} {site['sigungu']} {site['dong']} 반경 3km 생활권"),
            ("입지 종합 등급", f"{score['grade']}등급 (총점 {score['total_score']}점 / {score['grade_desc']})"),
            ("보고서 작성일", f"{data.get('created_at', '2026.08')}  |  마이파크 가맹본부 데이터전략실")
        ]
        for r_idx, (k, v) in enumerate(meta_data):
            c1 = meta_table.cell(r_idx, 0)
            c2 = meta_table.cell(r_idx, 1)
            c1.text = k
            c2.text = v
            c1.paragraphs[0].runs[0].font.bold = True
            c1.paragraphs[0].runs[0].font.size = Pt(9.5)
            c2.paragraphs[0].runs[0].font.size = Pt(9.5)
        
        doc.add_paragraph().paragraph_format.space_after = Pt(15)
        
        # 1. 입지 적합성 종합 판정
        h1 = doc.add_heading("1. 입지 적합성 종합 판정 (5-Dimension Diamond Scoring)", level=1)
        h1.runs[0].font.color.rgb = RGBColor(0, 43, 73)
        doc.add_paragraph(f"본 사업지는 5대 다이아몬드 스코어링 분석 결과 총점 {score['total_score']}점({score['grade']}등급 - {score['grade_desc']})으로 출점 최우선 추천 입지로 평가되었습니다.")
        doc.add_paragraph(f"• 시니어 인구 밀집도: {score['scores']['senior_population']}점 / 25점 (반경 3km 내 50대 이상 시니어 {demo['senior_50_plus']:,}명 밀집)")
        doc.add_paragraph(f"• 접근성 및 주차 인프라: {score['scores']['accessibility_parking']:.1f}점 / 25점 (간선도로 접면 및 대중교통 우수)")
        doc.add_paragraph(f"• 공간 적합성 및 층고: {score['scores']['space_efficiency']:.1f}점 / 15점 (전용 {site['area_pyeong']}평 10타석 배치 최적)")
        doc.add_paragraph(f"• 경쟁 매장 여유도: {score['scores']['supply_gap']}점 / 15점 ({comm.get('competitor_summary', '공급 절대 부족')})")
        doc.add_paragraph(f"• 지역 소비력 및 여가지출: {score['scores']['commercial_spending']}점 / 20점 (상위 20% 월매출 {comm['top_20_sales']//10000:,}만원 시장)")

        # 2. 배후 인구 및 시니어 타겟 분석
        h2 = doc.add_heading("2. 3km 생활권 인구 및 타겟 연령 분석", level=1)
        h2.runs[0].font.color.rgb = RGBColor(0, 43, 73)
        doc.add_paragraph(f"사업지 반경 3km(자동차 10분 생활권) 내 거주 총인구는 {demo['total_pop']:,}명(남 {demo['male_pop']:,}명 / 여 {demo['female_pop']:,}명)입니다. 이 중 스크린 파크골프의 핵심 타겟인 50대 이상 시니어 인구는 약 {demo['senior_50_plus']:,}명으로 전체 인구의 {demo['senior_ratio']}%를 차지합니다.")
        p_sub = doc.add_paragraph(f"특히 주간 시간대 소비를 주도하는 50대 이상 여성 인구가 {demo['senior_50_female']:,}명에 달해 평일 낮(10~17시) 주부 및 친목 동호회 정기 예약 유치에 최적의 환경을 형성하고 있습니다.")
        if demo.get('is_estimated'):
            p_sub.add_run(" (※ 본 인구 데이터는 행정동 추정치 모델이 적용되었으므로 현장 실측 조사를 병행 권장합니다.)")

        # 3. 상권 소비력 및 경쟁 환경
        h3 = doc.add_heading("3. 상권 매출 동향 및 경쟁 환경 분석", level=1)
        h3.runs[0].font.color.rgb = RGBColor(0, 43, 73)
        doc.add_paragraph(f"소상공인 365 및 BASA 빅데이터 분석 결과, 해당 권역 내 스포츠·골프 업종의 점포당 월평균 매출액은 약 {comm['monthly_avg_sales']//10000:,}만원(상위 20% 매장: {comm['top_20_sales']//10000:,}만원) 수준으로 안정적인 여가 소비력을 입증하고 있습니다.")
        doc.add_paragraph(f"• 시간대별 가동: 주간(10~17시) 이용 비중이 {comm['time_distribution']['주간_10_17시_비중']}%로 일반 스크린골프 대비 주간 가동률이 압도적으로 높습니다.")
        doc.add_paragraph(f"• 반경 3km 경쟁: 일반 스크린골프 대비 '10타석 규모 전문 스크린 파크골프' 시설이 전무하여 플래그십 선점 경쟁력이 탁월합니다.")

        # 4. 사업지 개요 및 현장 출점 요건
        h4 = doc.add_heading("4. 사업지 개요 및 현장 출점 요건 (건축·인프라 체크리스트)", level=1)
        h4.runs[0].font.color.rgb = RGBColor(0, 43, 73)
        doc.add_paragraph(f"본 사업지는 {site['full_address']}에 위치하며, 전용면적 약 {site['area_pyeong']}평 공간에 {site['rooms']}개의 마이파크 스크린 파크골프 타석 및 라운지/카페가 구축되는 쾌적한 플래그십 모델입니다.")
        doc.add_paragraph(f"• 층고 요건: {site['clear_height_spec']}")
        doc.add_paragraph(f"• 주차 요건: {site['parking_spec']}")
        doc.add_paragraph(f"• 이동 편의: {site['accessibility_spec']}")
        doc.add_paragraph(f"• 인허가 및 용도: {site['zoning_spec']}")
        if site.get('special_notes'):
            doc.add_paragraph(f"• 고객 특이사항 반영: {site['special_notes']}")

        # 5. 표준 투자 조건 및 사업 추진 유의사항 (신규)
        h5 = doc.add_heading("5. 표준 투자 조건 및 사업 추진 유의사항 (SSOT)", level=1)
        h5.runs[0].font.color.rgb = RGBColor(0, 43, 73)
        doc.add_paragraph("■ 표준 창업 투자비 (10타석 120평 플래그십 기준 SSOT):")
        doc.add_paragraph("• 시뮬레이터 장비(10대): 1억 5,000만원 (대당 1,500만원 고정)")
        doc.add_paragraph("• 인테리어 공사비(120평): 1억 4,400만원 (평당 120만원)")
        doc.add_paragraph("• 부대설비(냉난방/간판/가구/초도용품): 2,500만원")
        doc.add_paragraph(f"★ 총 초기 순투자비용: 총 {(inv['total_capex']/100000000.0):.2f}억원 ({(inv['total_capex']//10000):,}만원)")
        
        doc.add_paragraph("■ 사업 추진 필수 유의사항 (Caveat):")
        doc.add_paragraph("• 위 수치는 표준 모델 기준 추정치이며, 실제 임대료 및 인테리어 공사비는 현장 실측 견적에 따라 달라질 수 있습니다.")
        doc.add_paragraph("• 매니저/직원을 채용해 전면 위탁 운영할 경우 인건비 증가(월 500~750만원)로 손익분기점이 상승하고 회수기간이 늘어납니다.")
        doc.add_paragraph("• 본 보고서는 표준 재무 모델에 기반한 분석 자료이며, 실제 미래 사업 성과를 보장하지 않습니다.")

        # 6. 재무 타당성 및 투자금 회수 분석
        h6 = doc.add_heading("6. 재무 타당성 및 투자금 회수 분석 (3대 시나리오 & BEP)", level=1)
        h6.runs[0].font.color.rgb = RGBColor(0, 43, 73)
        
        c_sc = fin['monthly_scenarios']['conservative']
        m_sc = fin['monthly_scenarios']['moderate']
        o_sc = fin['monthly_scenarios']['optimistic']
        
        doc.add_paragraph("■ 3대 시나리오별 월간 예상 손익 (1인 18홀 7,000원 기준):")
        doc.add_paragraph(f"• 보수적 시나리오 (일 {c_sc['daily_turns_per_room']}회전): 월매출 {c_sc['total_revenue']//10000:,}만원 / 월 순영업이익 {c_sc['operating_profit']//10000:,}만원 (회수 {inv['payback_months_conservative']:.1f}개월)")
        doc.add_paragraph(f"• 보편적 시나리오 (일 {m_sc['daily_turns_per_room']}회전): 월매출 {m_sc['total_revenue']//10000:,}만원 / 월 순영업이익 {m_sc['operating_profit']//10000:,}만원 (회수 {inv['payback_months_moderate']:.1f}개월)")
        doc.add_paragraph(f"• 긍정적 시나리오 (일 {o_sc['daily_turns_per_room']}회전): 월매출 {o_sc['total_revenue']//10000:,}만원 / 월 순영업이익 {o_sc['operating_profit']//10000:,}만원 (회수 {inv['payback_months_optimistic']:.1f}개월)")
        
        doc.add_paragraph("■ 운영 방식에 따른 손익분기점(BEP) 및 회수 기간 비교:")
        doc.add_paragraph("1) 창업주 직접 상주 운영 모델 (점주 1인 상주, 인건비 월 250만원):")
        doc.add_paragraph(f"   • 월 고정비: 약 {fin['owner_operated']['fixed_cost']//10000:,}만원 (임대료 {site['monthly_rent']//10000:,}만 + 점주인건비 250만 + 관리비 230만)")
        doc.add_paragraph(f"   • 손익분기점(BEP): 타석당 하루 1팀 (일 {inv['bep_turns_per_room']}회전 / 1일 약 {inv['bep_daily_users']}명) 이용 시 월 고정비 전액 커버")
        doc.add_paragraph(f"   • 보편 가동 시 월 순영업이익: 약 {m_sc['operating_profit']//10000:,}만원 (영업이익률 {m_sc['profit_margin']}%)")
        doc.add_paragraph(f"   • 초기 투자금 3.19억원 전액 회수 기간: 단 {inv['payback_months_moderate']:.1f}개월 (약 1년 1개월 만에 원금 100% 전액 회수)")
        
        doc.add_paragraph("2) 직원 3명 채용 위탁 운영 모델 (매니저/직원 채용, 인건비 월 750만원):")
        doc.add_paragraph(f"   • 월 고정비: 약 {fin['owner_operated']['fixed_cost']//10000 + 500:,}만원, 손익분기점: 타석당 하루 1.6팀 (일 {fin['owner_operated']['staff3_bep_turns']}회전)")
        doc.add_paragraph(f"   • 보편 가동 시 월 순영업이익: 약 {fin['owner_operated']['staff3_operating_profit']//10000:,}만원, 회수 기간: 약 {fin['owner_operated']['staff3_payback_months']:.1f}개월")
        
        # 7. 종합 제언 및 기대효과
        h7 = doc.add_heading("7. 5개년 중장기 손익 전망 및 최종 종합 제언", level=1)
        h7.runs[0].font.color.rgb = RGBColor(0, 43, 73)
        doc.add_paragraph(f"5개년 누적 매출 약 {fin['five_year']['total_5yr_revenue']//100000000:.1f}억원, 누적 순영업이익 {fin['five_year']['total_5yr_profit']//100000000:.1f}억원이 예상되는 우수한 고수익·저위험 사업 모델입니다.")
        
        doc.add_paragraph("■ 가맹점 출점 기대효과 및 핵심 경쟁력:")
        doc.add_paragraph(score['value_franchisee'])
        
        doc.add_paragraph("■ 상가 전체 상권 활성화 및 건물 가치 상승 효과:")
        doc.add_paragraph(score['value_landlord'])
        
        os.makedirs(os.path.dirname(output_docx_path), exist_ok=True)
        doc.save(output_docx_path)
        print(f"[DOCX GENERATED] {output_docx_path}")
        return output_docx_path

DOCXGenerator = DocxGenerator

